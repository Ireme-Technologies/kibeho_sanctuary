#!/usr/bin/env python3
"""Build the shrine website content-request Word pack."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x12, 0x28, 0x47)
SKY = RGBColor(0x2A, 0x6F, 0x97)
MUTED = RGBColor(0x4A, 0x55, 0x63)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

OUT = Path(__file__).resolve().parent / "Kibeho-Shrine-Website-Content-Request.docx"


def set_run(run, size=11, bold=False, color=NAVY, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_para(doc, text, *, size=11, bold=False, color=NAVY, space_after=8, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13, bold=True, color=NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "2A6F97")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r2 = p.add_run(text)
        set_run(r2)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    set_run(r)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_run(run, size=10, bold=True, color=WHITE)
        shade_cell(cell, "122847")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            set_run(run, size=10, color=NAVY)
            if r_i % 2 == 1:
                shade_cell(cell, "F4F1EA")
    doc.add_paragraph()
    return t


def blank_rows(n, cols):
    return [[" "] * cols for _ in range(n)]


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "Shrine of Our Lady of Kibeho  ·  Official website content request  ·  For internal use"
    )
    set_run(fr, size=8, color=MUTED)

    add_para(doc, "DIOCESE OF GIKONGORO", size=10, bold=True, color=SKY, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Content request for the official website",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "Shrine of Our Lady of Kibeho  ·  Sanctuaire Notre-Dame de Kibeho  ·  Umwibutso wa Nyina wa Kibeho",
        size=11,
        italic=True,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_para(
        doc,
        "Prepared for the Rector, the Pilgrimage Office, and the Diocese  ·  Please return with attachments",
        size=10,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )

    heading(doc, "1. Why this pack exists")
    add_para(
        doc,
        "The website structure is ready: Our Lady, the Shrine, Pilgrimage, Spirituality, News, and Support. "
        "What it still needs is the Shrine’s own voice — names, times, prayers, photos, and facts that only "
        "the sanctuary can confirm. Placeholder text and sample events cannot be the public face of Africa’s "
        "Church-recognised Marian apparition.",
    )
    add_para(
        doc,
        "This is not a survey of opinions. It is a request for materials we will publish, with your approval. "
        "Short, exact answers are better than long essays. Official documents and photographs are better than "
        "retyping.",
    )

    heading(doc, "2. The best way to reply (please read this first)")
    add_para(
        doc,
        "Do not put everything into a Google Form. Forms are excellent for telephone numbers and yes/no facts. "
        "They are a poor place for the 2001 declaration, the official prayer, Mass tables, bank letters, and photo albums.",
        italic=True,
    )
    bullet(doc, " Upload this file to Google Drive, convert it to a Google Doc, and share it with the people named in section 3. They can type in the tables and comment in the margin.", bold_lead="Preferred. ")
    bullet(doc, " Keep this Word file, fill the tables, and return it by email with a zip of attachments.", bold_lead="Also good. ")
    bullet(doc, " Use a Google Form only for Pack 2’s directory facts (phones, Mass times grid) if the office prefers a short form. A Form cannot replace Packs 1, 3, 4 or 5.", bold_lead="Optional extra. ")
    add_para(
        doc,
        "Answer in the language that is most accurate — Kinyarwanda or French is welcome. We will prepare English (and German where needed). "
        "Please mark any sentence that must not be changed (official titles, the wording of the recognition, prayers).",
    )
    table(
        doc,
        ["Coordinator (name, role, WhatsApp, email)", "Target return date"],
        [[" ", " "]],
    )

    heading(doc, "3. Who fills which pack")
    table(
        doc,
        ["Pack", "Who should answer", "Needed for launch?"],
        [
            ["1. Identity and approval", "Rector / Bishop’s office", "Yes — first"],
            ["2. Coming on pilgrimage", "Pilgrimage Office", "Yes — first"],
            ["3. Message, prayers, visionaries", "Rector / chaplain", "Yes — first"],
            ["4. Places, map, photos", "Pilgrimage Office + communications", "Yes — as soon as photos exist"],
            ["5. Support, Master Plan, money", "Diocesan finance / development", "Yes for donations; figures can follow"],
            ["6. News, video, languages", "Communications", "Can start after Packs 1–3"],
        ],
    )

    heading(doc, "4. What we already have (please correct, do not ignore)")
    add_para(doc, "We will not guess over your corrections. Strike through anything that is wrong.")
    bullet(doc, "Official name to use in public: Shrine of Our Lady of Kibeho (not “Kibeho Sanctuary” as the main title). Confirm or give the exact form in EN / FR / RW / DE.")
    bullet(doc, "Address on file: Sanctuary Our Lady of Kibeho, B.P. 341 Butare / Rwanda. Plus Code 9H23+58 Kibeho.")
    bullet(doc, "Phones on file: +250 788 559 192 and +250 788 307 376. Email: info@kibehosanctuary.rw.")
    bullet(doc, "Facebook NyinaWaJambo · Instagram sanctuairenotredamedekibeho · X kibehosanctuary.")
    bullet(doc, "Recognition date in our draft: 29 June 2001, Bishop Augustin Misago of Gikongoro, with Holy See approval. Three recognised visionaries: Alphonsine Mumureke, Nathalie Mukamazimpaka, Marie Claire Mukangango.")
    bullet(doc, "Bank names and numbers currently published for donations (BK RWF/EUR/USD and BPR RWF, plus a MoMo code). These must be confirmed or replaced before the site stays public.")
    bullet(doc, "Candle offering USD 1 · Mass offering USD 20 — confirm or give the real amounts and currencies.")

    heading(doc, "PACK 1 — Identity and approval")
    add_para(doc, "Owner: Rector / Bishop’s office. Goal: one authorised public identity.", italic=True)
    numbered(doc, "Exact official name in Kinyarwanda, French, English, and German (if used). Any short name allowed in the logo?")
    numbered(doc, "Who is the legal publisher of the website (Diocese of Gikongoro / Shrine / both)?")
    numbered(doc, "Who may approve texts before they go live (name and role)? Who may approve photographs of liturgy, pilgrims, and minors?")
    numbered(doc, "Preferred public email, two phones, WhatsApp for pilgrims, and postal address. Any second office in Kigali (e.g. Pallotti Pilgrimages Centre) — name, address, phone, what they actually do?")
    numbered(doc, "Social accounts that are official. Any that must not be linked?")
    numbered(doc, "Pastoral team to list publicly. For each person: full name, title (Rector, chaplain, …), two-sentence biography, photo we may use, whether to show email.")
    table(
        doc,
        ["Full name", "Title / role", "May we publish a photo?", "Public contact?"],
        blank_rows(6, 4),
    )
    numbered(doc, "Religious communities around the Shrine that should appear. For each: official name, what they do at Kibeho, website if any, photo.")
    table(
        doc,
        ["Community (official name)", "Role at the Shrine", "Contact (optional)"],
        blank_rows(8, 3),
    )

    heading(doc, "PACK 2 — Coming on pilgrimage")
    add_para(doc, "Owner: Pilgrimage Office. Goal: a pilgrim can plan a real visit.", italic=True)
    numbered(doc, "How should a parish, a family, or a foreign group register? What must they send (dates, number of people, language, priest travelling with them)?")
    numbered(doc, "Typical reply time of the office. Opening hours.")
    numbered(doc, "Ordinary weekday and Sunday Mass: day, title (e.g. Morning Mass), start time, end time, language, place (which church/chapel), notes (confession before Mass, etc.). Mark feast-day exceptions.")
    table(
        doc,
        ["Day", "Celebration", "Start", "End", "Language", "Place", "Notes"],
        blank_rows(10, 7),
    )
    numbered(doc, "Confession, Adoration, communal Rosary / Seven Sorrows: usual times, or “as announced”.")
    numbered(doc, "Annual calendar for the next 12–24 months. For each public gathering: official title, type (feast / national pilgrimage / international / youth / retreat), start and end dates, start and end times, whether it repeats every year, place, 5–8 line description, and whether groups must register.")
    table(
        doc,
        ["Title", "Type", "Dates", "Times", "Annual?", "Register?"],
        blank_rows(8, 6),
    )
    numbered(doc, "How to reach Kibeho from Kigali and from Huye today (recommended route, typical duration, any bus or hired-vehicle advice). Confirm: Kigali–Huye–Matyazo–Kibeho and other routes we may publish.")
    numbered(doc, "Practical rules we must publish: dress, photography, silence, offerings, water from the spring, what pilgrims should not do.")
    numbered(doc, "Accommodation the Shrine endorses. For each place: official name (spelling), type (guest house / hotel / apartments / retreat house), who manages it, walking time to the Shrine, phone, WhatsApp, email, website or booking page, 3-line description, 3–6 photos. Do not use hotel star ratings unless you insist.")
    table(
        doc,
        ["Name (correct spelling)", "Type", "Distance / walk", "Phone / WhatsApp", "Website", "Shrine-run?"],
        blank_rows(6, 6),
    )

    heading(doc, "PACK 3 — The message, the visionaries, the prayers")
    add_para(doc, "Owner: Rector / chaplain. Goal: a story the Church can stand behind.", italic=True)
    numbered(doc, "Please attach or point us to the texts we may quote: the 29 June 2001 declaration (or an official summary), any Holy See note you wish cited, and a short approved chronology of 1981–1989.")
    numbered(doc, "For each recognised visionary: names we may use today, dates of apparitions, what the Church associates with her (e.g. Seven Sorrows and Marie Claire), whether she may be photographed, any sentence she must not be reduced to.")
    numbered(doc, "The heart of the message, in the words you want on the website (conversion, prayer, Rosary of the Seven Sorrows, reparation, reconciliation). Any quote that is authorised.")
    numbered(doc, "Official prayer(s) of the Shrine — full text in RW / FR / EN if they exist. Who composed them? May we publish them?")
    numbered(doc, "How you teach pilgrims to pray the Seven Sorrows Rosary here (order of the seven sorrows, beads, any Kibeho-specific instruction). A novena if there is an official one.")
    numbered(doc, "Candle and Mass intentions: real price, currency, what the pilgrim must write, how payment is confirmed, any text you want on those pages.")
    numbered(doc, "Testimonies: may we publish named stories? What approval process? Two or three testimonies you already bless, or a rule that we only take future submissions through the office.")

    heading(doc, "PACK 4 — Places, map, and photographs")
    add_para(doc, "Owner: Pilgrimage Office and whoever holds the photo archive.", italic=True)
    numbered(doc, "List every church, chapel, and apparition place a pilgrim should know, with the name you want on the map. For each: 4–8 line description, any historical date, a photograph we may use.")
    table(
        doc,
        ["Place name", "Type (church / chapel / site / spring / stations)", "One-sentence description", "Photo filename"],
        [
            ["Church of Our Lady of Sorrows", "Church", " ", " "],
            ["Chapel of the Apparitions", "Chapel", " ", " "],
            ["Place of the apparitions", "Site", " ", " "],
            ["Esplanade of the apparitions", "Site", " ", " "],
            ["Chapel of Adoration", "Chapel", " ", " "],
            ["Source of Mary / Holy Spring", "Spring", " ", " "],
            ["Way of the Cross", "Path", " ", " "],
            ["Way of the Rosary / Seven Sorrows (if distinct)", "Path", " ", " "],
            [" ", " ", " ", " "],
            [" ", " ", " ", " "],
        ],
    )
    numbered(doc, "A recommended order of visit for a first-time pilgrim (numbered 1–7). Where should groups report first?")
    numbered(doc, "Photographs we need (landscape, high resolution, you hold the rights, no unconsented close-ups of children): hillside and churches; liturgy; pilgrims at prayer; spring; Way of the Cross; feast day; quiet weekday. For each file: date, photographer, caption, permission to use on web and social media.")
    numbered(doc, "Any site plan, sketch, or drone view we may redraw as a simple shrine map. If none, a hand-drawn sketch is enough.")
    numbered(doc, "One YouTube (or similar) video that may represent the Shrine on the site. Title and what it is (message, liturgy, documentary).")

    heading(doc, "PACK 5 — Support, Master Plan, transparency")
    add_para(doc, "Owner: Diocesan finance / development, with Rector. Goal: partners and donors trust the page.", italic=True)
    numbered(doc, "Confirm or replace every bank account and MoMo code. Account name exactly as it must appear. Any instruction pilgrims must put in the payment reference.")
    table(
        doc,
        ["Bank", "Account name (exact)", "Number", "Currency", "Still valid? (yes/no)"],
        [
            ["Bank of Kigali", "Diocese Gikongoro / Sanct KIBEHO", "00266 00690793-01", "RWF", " "],
            ["Bank of Kigali", "Diocese Gikongoro / Sanct KIBEHO", "00266 00690796-02", "EUR", " "],
            ["Bank of Kigali", "Diocese Gikongoro / Sanct KIBEHO", "00266 00690797-03", "USD", " "],
            ["BPR", "Diocese Gikongoro / Sanct KIBEHO", "475453520910197", "RWF", " "],
            ["MoMo", " ", "*182*8*1*060974#", "RWF", " "],
        ],
    )
    numbered(doc, "Master Plan: is there an approved document we may summarise? Phase names, what is built, what is next, what a partner would fund. Attach a PDF even if it is not for the public — we will only publish what you mark “public”.")
    numbered(doc, "Current projects for the website. For each: title, status (planning / in progress / completed), the need, what you will do, fruit for the local community, for the Church, for pilgrims from abroad, photos of the site or a drawing, whether you will publish a funding goal.")
    table(
        doc,
        ["Project title", "Status / phase", "The need (2–4 lines)", "May we show a funding goal?"],
        [
            ["Master Plan — Phase One", " ", " ", " "],
            ["Pilgrim Welcome Centre", " ", " ", " "],
            [" ", " ", " ", " "],
            [" ", " ", " ", " "],
        ],
    )
    numbered(doc, "Annual report or financial summary you are willing to publish (year, PDF). If none yet, a one-paragraph stewardship statement the Bishop or Rector will sign.")
    numbered(doc, "Partners who may be named and logo-listed (Diocese, congregations, “friends of Kibeho”, international groups). Any that must not appear.")

    heading(doc, "PACK 6 — News, tone, and languages")
    add_para(doc, "Owner: whoever writes for the Shrine publicly.", italic=True)
    numbered(doc, "Who sends news to the website (role, email)? How often?")
    numbered(doc, "Categories you want: News, Events, Rector, Bishop, Press — keep, drop, or rename.")
    numbered(doc, "Launch languages in order of importance: Kinyarwanda, French, English, German. Who will review each language?")
    numbered(doc, "Words we must never use. Words you prefer (e.g. Shrine not Sanctuary in English; Nyina wa Jambo).")
    numbered(doc, "Any existing leaflet, pilgrim booklet, or FAQ we should treat as the source of truth.")

    heading(doc, "5. Attachment checklist")
    add_para(doc, "Return a folder (Google Drive or zip) named with the date. Suggested file names in English or French are fine.")
    table(
        doc,
        ["File", "Have it? (yes / later)", "Notes"],
        [
            ["2001 recognition text or official summary", " ", " "],
            ["Official prayer(s) RW / FR / EN", " ", " "],
            ["Seven Sorrows how-to (if written)", " ", " "],
            ["Mass timetable (even a photo of a noticeboard)", " ", " "],
            ["Calendar of feasts and pilgrimages", " ", " "],
            ["Bank letter or stamped account list", " ", " "],
            ["Master Plan PDF (mark public / internal)", " ", " "],
            ["Pastoral team photos", " ", " "],
            ["Place photos (churches, spring, stations)", " ", " "],
            ["Hero / homepage photograph", " ", " "],
            ["Logo files (vector if possible)", " ", " "],
            ["One authorised video link", " ", " "],
            ["Site sketch or plan", " ", " "],
            ["Existing pilgrim leaflet / FAQ", " ", " "],
        ],
    )

    heading(doc, "6. If you still want a Google Form")
    add_para(
        doc,
        "A Form is useful as a short companion, not as the whole request. If you create one, keep it to directory facts so the office can finish it on a phone:",
    )
    bullet(doc, "Confirm name, phones, email, WhatsApp, address.")
    bullet(doc, "Office hours and how groups register.")
    bullet(doc, "Mass grid (or “we will send a photo of the timetable”).")
    bullet(doc, "Lodging names and phones.")
    bullet(doc, "Upload: one zip or a Drive link — Forms file upload is limited.")
    add_para(
        doc,
        "Put the long story, prayers, recognition text, Master Plan, and photo captions in this document (or in the Drive folder). Mixing both in one Form usually means the Rector never finishes it.",
        italic=True,
    )

    heading(doc, "7. After you send this")
    add_para(
        doc,
        "We will type approved answers into the website, send you a preview of each pillar (Our Lady, Shrine, Pilgrimage, Spirituality, News, Support), and not publish a page you have marked as internal. Incomplete packs are still useful: send Packs 1–3 first if the rest will take time.",
    )
    add_para(doc, "Thank you for helping the Shrine speak with its own voice.", italic=True, space_after=2)
    add_para(doc, "Ireme Tech  ·  for the Shrine of Our Lady of Kibeho, Diocese of Gikongoro", size=10, color=MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
